# -*- coding: utf-8 -*-
"""
חילוץ רישיונות יבואן (docx) לאקסל מנורמל — יבואן ישיר / עקיף / זעיר.
שימוש: python extract_licenses.py <תיקיית_קבצים> <פלט.xlsx>

קבצים נלווים (באותה תיקייה של הסקריפט):
  classifications.csv — טבלת תרגום סיווגים (מקור,יעד). ניתן להרחיב.

פלט:
  <פלט.xlsx>  — גיליונות: ישיר-חברות, ישיר-תוצרים, עקיף-חברות, עקיף-תוצרים, זעיר-חברות
                כל גיליון חברות כולל עמודת "סוג ישות" (תאגיד רשום / עוסק מורשה),
                מזוהה לפי "ח.פ" או "עוסק מורשה" במסמך — שני המספרים נכנסים לעמודת ח.פ.
                עקיף-תוצרים כולל תאריכי תוקף נפרדים לסוכן ראשי ולסוכן משני.
  extraction.log — אזהרות ושגיאות
  unknown_classifications.csv — סיווגים שלא נמצאו בטבלת התרגום (ערך, מופעים, קובץ לדוגמה)
"""
import csv
import logging
import re
import sys
from collections import Counter
from pathlib import Path

import pandas as pd
from docx import Document

HP_RE = re.compile(r"ח\.?\s?פ\.?\s*(\d{9})")
OSEK_RE = re.compile(r"עוסק(?:\s+מורשה)?[^\d]{0,20}(\d{9})")
LICENSE_VALID_RE = re.compile(r"תוקף\s+הרישיון\s+עד\s+ל(?:תאריך|יום)\s*:?\s*(\d{1,2}[./]\d{1,2}[./]\d{2,4})")
APPENDIX_DATE_RE = re.compile(r"מעודכן\s+ל(?:תאריך|יום)\s*:?\s*(\d{1,2}[./]\d{1,2}[./]\d{2,4})")
COMPANY_RE = re.compile(r"לחברת\s+(.+)")
INDIVIDUAL_RE = re.compile(r"ל(?:מר|גברת|גב['׳]?)\s+(.+)")
GRANT_DATE_RE = re.compile(r"היום הזה\s+(.+?)\s*\.?\s*$")
PARENS_RE = re.compile(r"^(.*?)\s*\(([^)]*)\)\s*$")
AGENT_ROLE_RE = re.compile(r"^(.*?)\s*\(\s*סוכן\s+(ראשי|משני|משנה)\s*\)\s*$")
DMY_RE = re.compile(r"^(\d{1,2})[./](\d{1,2})[./](\d{2,4})$")
LICENSE_TYPES = ["יבואן ישיר", "יבואן עקיף", "יבואן זעיר"]

log = logging.getLogger("extract")
unknown_class = Counter()
unknown_example = {}


def load_translation(script_dir: Path) -> dict:
    path = script_dir / "classifications.csv"
    table = {}
    with open(path, encoding="utf-8-sig") as f:
        for row in csv.DictReader(f):
            key = " ".join(row["מקור"].split())
            table[key] = row["יעד"].strip()
    return table


def parse_date(raw, source):
    raw = (raw or "").strip()
    if not raw:
        return ""
    m = DMY_RE.match(raw)
    if m:
        d, mo, y = map(int, m.groups())
        if y < 100:
            y += 2000
        try:
            return pd.Timestamp(year=y, month=mo, day=d)
        except ValueError:
            pass
    log.warning("%s: תאריך לא תקין '%s' — נשמר כטקסט", source, raw)
    return raw


CODE_RUN_RE = re.compile(r"^(?:L7|[MNO][1-4]|L)+$")
CODE_TOKEN_RE = re.compile(r"L7|[MNO][1-4]|L")


def expand_code_token(tok: str, table: dict):
    """'N1M2' -> ['N1','M2'] (קודים שנדבקו זה לזה בטעות הקלדה במקור).
    מוחזר None אם הטוקן אינו רצף קודים תקין."""
    if tok in table or len(tok) < 4:
        return None
    if not CODE_RUN_RE.match(tok):
        return None
    parts = CODE_TOKEN_RE.findall(tok)
    return parts if len(parts) >= 2 else None


def split_product_cell(cell: str, table: dict):
    """מפריד תא 'תוצר' לשם התוצר ולרשימת קודי הסיווג שלו. תומך בכמה פורמטים:
    'פולקסווגן (M1,M2,N1)'      -> ('פולקסווגן', ['M1','M2','N1'])   [סוגריים]
    'טויוטה M1, M2, N1'         -> ('טויוטה', ['M1','M2','N1'])      [פסיקים בסוף, בלי סוגריים]
    'טויוטה M1 N1 M2 N2'        -> ('טויוטה', ['M1','N1','M2','N2']) [רווחים בלבד, בלי פסיקים/סוגריים]
    'LOADRITE O1'                -> ('LOADRITE', ['O1'])              [טוקן בודד מוכר בסוף]
    """
    cell = " ".join(cell.split())
    # סוגריים לא מאוזנים במקור (למשל 'DONGFENG M1 M2))' או 'איווקו ((N1') —
    # מסירים את כל תווי הסוגריים ומטפלים בשאר לפי מפרידים רגילים
    if cell.count("(") != cell.count(")"):
        cell = " ".join(cell.replace("(", " ").replace(")", " ").split())
    m = PARENS_RE.match(cell)
    if m:
        name = m.group(1).strip()
        inner = m.group(2).strip()
        if inner in table:
            # ערך רב-מילים מוכר כמכלול (למשל 'תלת אופנוע') — לא מפצלים
            items = [inner]
        else:
            items = []
            for part in inner.split(","):
                part = " ".join(part.split())
                if not part:
                    continue
                if part in table:
                    items.append(part)  # ערך מוכר (כולל רב-מילים) — נשמר שלם
                else:
                    # ערך לא-מוכר עם רווחים (למשל 'N2 M2') — מפצלים לקודים
                    items.extend(part.split())
        return name, items

    # בלי סוגריים, עם פסיקים: קולפים מהסוף רצף רציף של קודים מוכרים בטבלת התרגום
    if "," in cell:
        parts = [p.strip() for p in cell.split(",")]
        idx = len(parts)
        codes = []
        while idx > 0 and parts[idx - 1] in table:
            codes.insert(0, parts[idx - 1])
            idx -= 1
        # גבול: הקוד הראשון עלול להיות דבוק לשם בלי פסיק, למשל 'טויוטה M1, M2'
        if idx > 0:
            bwords = parts[idx - 1].split()
            if len(bwords) > 1 and bwords[-1] in table:
                codes.insert(0, bwords[-1])
                parts[idx - 1] = " ".join(bwords[:-1])
        if codes:
            name = ", ".join(p for p in parts[:idx] if p).strip()
            return name or cell, codes

    # בלי סוגריים ובלי פסיקים: קולפים מהסוף מילים בודדות שכל אחת בנפרד קוד מוכר
    words = cell.split()
    idx = len(words)
    codes = []
    while idx > 0:
        tok = words[idx - 1]
        if tok in table:
            codes.insert(0, tok)
        else:
            expanded = expand_code_token(tok, table)
            if not expanded:
                break
            codes[0:0] = expanded
        idx -= 1
    if codes:
        return " ".join(words[:idx]).strip() or cell, codes

    # בלי סוגריים ובלי פסיקים: טוקן/ים בודדים מוכרים בסוף (מופרדים ברווח)
    words = cell.split()
    for n in (3, 2, 1):
        if len(words) > n:
            tail = " ".join(words[-n:])
            if tail in table:
                return " ".join(words[:-n]), [tail]
    return cell, []


def split_agent_cell(cell: str):
    """'HAO INVESTMENT GROUP LLC (סוכן ראשי)' -> ('HAO INVESTMENT GROUP LLC', 'ראשי')
    אם התפקיד לא מזוהה בתא — מוחזר התא כמות שהוא ותפקיד ריק."""
    cell = " ".join(cell.split())
    m = AGENT_ROLE_RE.match(cell)
    if m:
        role = "משני" if m.group(2) == "משנה" else m.group(2)
        return m.group(1).strip(), role
    return cell, ""


def translate(items: list, table: dict, fname: str):
    """מחזיר (target_list, all_found: bool, changed: bool)."""
    targets, all_found, changed = [], True, False
    for it in items:
        key = " ".join(it.split())
        if key in table:
            targets.append(table[key])
            if table[key] != key:
                changed = True
            continue
        # 'O3 O4' / 'M1 N1 N2' — כמה קודים מופרדים ברווח בלי פסיק
        tokens = key.split()
        if len(tokens) > 1 and all(t in table for t in tokens):
            targets.extend(table[t] for t in tokens)
            changed = True
            continue
        targets.append(key)
        all_found = False
        unknown_class[key] += 1
        unknown_example.setdefault(key, fname)
    return targets, all_found, changed


def find_manager(doc) -> str:
    paras = [p.text.strip() for p in doc.paragraphs]
    for i, t in enumerate(paras):
        if t.replace(" ", "") == "המנהל":
            for j in range(i - 1, -1, -1):
                cand = paras[j].strip()
                if not cand or set(cand) <= {"_", " "}:
                    continue
                if "מנהל" in cand:  # שורת תפקיד/כותרת (למשל 'מנהל אגף...'), לא שם
                    continue
                return " ".join(cand.split())
    return ""


def dedup_cells(row):
    """מחזיר (טקסטים, אלמנטים) של תאי השורה לאחר קיפול תאים ממוזגים אופקית.
    python-docx חוזר על אותו תא עבור טווח ממוזג — מזהים לפי אלמנט ה-tc."""
    texts, elems, seen = [], [], set()
    for c in row.cells:
        el = c._tc
        if id(el) in seen:
            continue
        seen.add(id(el))
        texts.append(" ".join(c.text.split()))
        elems.append(el)
    return texts, elems


def extract_file(path: Path, table: dict):
    doc = Document(path)
    paras = [p.text for p in doc.paragraphs]
    full_text = "\n".join(paras)

    hp_m = HP_RE.search(full_text)
    osek_m = OSEK_RE.search(full_text)
    if hp_m:
        hp = hp_m.group(1)
        entity_type = "תאגיד רשום"
    elif osek_m:
        hp = osek_m.group(1)
        entity_type = "עוסק מורשה"
    else:
        hp = ""
        entity_type = ""
    comp_m = COMPANY_RE.search(full_text)
    if comp_m:
        company_name = comp_m.group(1).strip()
    else:
        ind_m = INDIVIDUAL_RE.search(full_text)
        company_name = ind_m.group(1).strip() if ind_m else ""
    license_type = next((t for t in LICENSE_TYPES if t in full_text), "")
    grant_date = ""
    for p in paras:
        gm = GRANT_DATE_RE.search(" ".join(p.split()))
        if gm and "לחודש" in gm.group(1):
            grant_date = gm.group(1)
            break
    manager = find_manager(doc)
    update_dates = APPENDIX_DATE_RE.findall(full_text)
    valid_m = LICENSE_VALID_RE.search(full_text)
    license_valid_date = parse_date(valid_m.group(1), path.name) if valid_m else ""

    cancelled = "מבוטל" in full_text or "בוטל" in full_text
    cancel_note = ""
    if cancelled:
        cancel_note = next((" ".join(p.split()) for p in paras if "בוטל" in p), "")

    missing = [n for n, v in [("ח.פ / עוסק מורשה", hp), ("שם חברה", company_name),
                              ("סוג רישיון", license_type)] if not v]
    for f in missing:
        log.warning("%s: לא נמצא שדה '%s'", path.name, f)

    products, prod_ok = [], True
    prev_col = None  # מיפוי עמודות מהטבלה הקודמת — לשימוש בטבלאות המשך ללא כותרת
    for t_idx, tbl in enumerate(doc.tables):
        if not tbl.rows:
            continue
        header, _ = dedup_cells(tbl.rows[0])
        has_header = any("תוצר" in h for h in header)

        if has_header:
            col = {}
            for i, h in enumerate(header):
                if "תוצר" in h:
                    col["product"] = i
                elif "יצרן" in h:
                    col["manufacturer"] = i
                elif "תוקף" in h and "ראשי" in h:
                    col["valid_agent1"] = i
                elif "תוקף" in h and ("משני" in h or "משנה" in h):
                    col["valid_agent2"] = i
                elif "תוקף" in h:
                    col["valid"] = i
                elif "תנאים" in h or "הגבלות" in h:
                    col["conditions"] = i
                elif "סוכן" in h and "ראשי" in h:
                    col["agent1"] = i
                elif "סוכן" in h and ("משני" in h or "משנה" in h):
                    col["agent2"] = i
                elif "סוכן" in h or "דילר" in h:
                    # עמודת סוכן/דילר גנרית — התפקיד ראשי/משני מוטמע בתוך תא הערך
                    # עצמו כ-'(סוכן ראשי)'/'(סוכן משני)' וייפרס לשורה
                    col["agent_generic"] = i
            prev_col = col
            data_rows = tbl.rows[1:]
        else:
            # טבלת המשך ללא כותרת (פיצול של נספח לאותו מבנה) — משתמשים במיפוי הקודם
            # וכל השורות הן נתונים (כולל הראשונה)
            if prev_col is None:
                continue
            col = prev_col
            data_rows = tbl.rows

        prev_product_el = None  # אלמנט תא-התוצר של השורה הקודמת, לזיהוי מיזוג אנכי
        for row in data_rows:
            cells, elems = dedup_cells(row)
            if not any(cells):
                continue

            def get(key):
                return cells[col[key]] if key in col and col[key] < len(cells) else ""

            raw_product = get("product")
            generic_valid = parse_date(get("valid"), path.name) if "valid" in col else ""
            product_el = elems[col["product"]] if "product" in col and col["product"] < len(elems) else None

            # שורת המשך (סוכן משני): מזוהה אם תא התוצר ממוזג אנכית עם השורה הקודמת
            # (אותו אלמנט tc) או שהתוצר ריק — ובשני המקרים יש סוכן והתוצר לא חדש
            is_continuation = (
                "agent_generic" in col and products and
                ((product_el is not None and product_el is prev_product_el) or not raw_product)
            )
            if is_continuation:
                agent_name, role = split_agent_cell(get("agent_generic"))
                if agent_name:
                    prev = products[-1]
                    if role == "משני" or (not role and prev["סוכן משני"] == ""):
                        if not role:
                            log.warning("%s: שורת המשך ללא תפקיד עבור '%s' — שויך לסוכן משני",
                                        path.name, agent_name)
                        prev["סוכן משני"] = agent_name
                        prev["תאריך תוקף סוכן משני"] = generic_valid
                    else:
                        prev["סוכן ראשי"] = agent_name
                        prev["תאריך תוקף סוכן ראשי"] = generic_valid
                continue

            prev_product_el = product_el

            name, src_items = split_product_cell(raw_product, table)
            tgt_items, found, changed = translate(src_items, table, path.name)
            if not found:
                prod_ok = False

            agent1_name = get("agent1")
            agent2_name = get("agent2")
            valid_agent1 = parse_date(get("valid_agent1"), path.name) if "valid_agent1" in col else ""
            valid_agent2 = parse_date(get("valid_agent2"), path.name) if "valid_agent2" in col else ""

            if "agent_generic" in col:
                # עמודת סוכן/דילר יחידה: השורה נושאת תוצר, לכן הסוכן שלה ראשי כברירת מחדל
                agent_name, role = split_agent_cell(get("agent_generic"))
                if role == "משני":
                    agent2_name, valid_agent2 = agent_name, generic_valid
                else:
                    agent1_name, valid_agent1 = agent_name, generic_valid

            products.append({
                "ח.פ": hp,
                "תוצר": name,
                "יצרן": get("manufacturer"),
                "רשימת סיווגים - מקור": ", ".join(src_items),
                "רשימת סיווגים - יעד": ", ".join(tgt_items),
                "סוכן ראשי": agent1_name,
                "סוכן משני": agent2_name,
                "תאריך תוקף": generic_valid if "agent_generic" not in col else "",
                "תאריך תוקף סוכן ראשי": valid_agent1,
                "תאריך תוקף סוכן משני": valid_agent2,
                "תנאים והגבלות": get("conditions"),
                "האם בוצעה המרה של סיווגים": "כן" if changed else "לא",
                "סטטוס OCR": "הצליח" if found and raw_product else "חלקי",
            })

    status = "הצליח"
    if missing or not prod_ok:
        status = "חלקי"
    if license_type != "יבואן זעיר" and not products:
        log.warning("%s: לא נמצאה טבלת נספח", path.name)
        status = "חלקי"

    company = {
        "ח.פ": hp,
        "סוג ישות": entity_type,
        "שם חברה": company_name,
        "סוג רישיון": license_type,
        "תאריך תוקף הרישיון": license_valid_date,
        "תאריך מתן הרישיון": grant_date,
        "שם המנהל המאשר של הרישיון": manager,
        "תאריך עדכון אחרון נספח א": parse_date(update_dates[-1], path.name) if update_dates else "",
        "האם הרישיון מבוטל ?": "כן" if cancelled else "לא",
        "הערה בנושא ביטול הרישיון": cancel_note,
        "שם קובץ מקור": path.name,
        "סטטוס OCR": status,
    }
    return license_type, company, products


DIRECT_COMP_COLS = ["ח.פ", "סוג ישות", "שם חברה", "סוג רישיון", "תאריך מתן הרישיון",
                    "שם המנהל המאשר של הרישיון", "תאריך עדכון אחרון נספח א",
                    "שם קובץ מקור", "סטטוס OCR"]
INDIRECT_COMP_COLS = ["ח.פ", "סוג ישות", "שם חברה", "סוג רישיון", "תאריך עדכון אחרון נספח א",
                      "תאריך מתן הרישיון", "שם המנהל המאשר של הרישיון",
                      "שם קובץ מקור", "סטטוס OCR"]
ZAIR_COMP_COLS = ["ח.פ", "סוג ישות", "שם חברה", "סוג רישיון", "תאריך תוקף הרישיון",
                  "שם קובץ מקור", "האם הרישיון מבוטל ?", "הערה בנושא ביטול הרישיון", "סטטוס OCR"]
DIRECT_PROD_COLS = ["ח.פ", "תוצר", "יצרן", "רשימת סיווגים - מקור",
                    "רשימת סיווגים - יעד", "תאריך תוקף", "תנאים והגבלות",
                    "האם בוצעה המרה של סיווגים", "סטטוס OCR"]
INDIRECT_PROD_COLS = ["ח.פ", "תוצר", "יצרן", "רשימת סיווגים - מקור",
                      "רשימת סיווגים - יעד", "סוכן ראשי", "סוכן משני",
                      "תאריך תוקף סוכן ראשי", "תאריך תוקף סוכן משני", "תנאים והגבלות",
                      "האם בוצעה המרה של סיווגים", "סטטוס OCR"]


def build_distinct_products(buckets: dict) -> pd.DataFrame:
    """מרכז רשימת תוצרים ייחודית (DISTINCT) מכל סוגי הרישיונות שיש בהם תוצרים.
    מפתח הייחודיות הוא שם התוצר המנורמל (trim + רווחים מקופלים), אך שומרים את
    הכתיב המקורי הראשון שנראה. לכל תוצר: היצרנים שנצפו, מספר מופעים כולל,
    וסוגי הרישיון שבהם הופיע — כדי לסייע בזיהוי כפילויות עברית/אנגלית וקיצורים."""
    agg = {}  # norm_name -> dict
    for ltype, (_, prods) in buckets.items():
        for p in prods:
            name = (p.get("תוצר") or "").strip()
            if not name:
                continue
            key = " ".join(name.split())
            rec = agg.setdefault(key, {
                "תוצר": name, "יצרנים": {}, "מופעים": 0, "סוגי רישיון": set()})
            rec["מופעים"] += 1
            rec["סוגי רישיון"].add(ltype.replace("יבואן ", ""))
            man = (p.get("יצרן") or "").strip()
            if man:
                rec["יצרנים"][man] = rec["יצרנים"].get(man, 0) + 1

    rows = []
    for rec in agg.values():
        # יצרנים ממוינים לפי שכיחות (הנפוץ ראשון) — עוזר בזיהוי מיפוי כפילויות
        mans = sorted(rec["יצרנים"].items(), key=lambda kv: (-kv[1], kv[0]))
        rows.append({
            "תוצר": rec["תוצר"],
            "יצרן": " | ".join(m for m, _ in mans),
            "מספר מופעים": rec["מופעים"],
            "סוגי רישיון": ", ".join(sorted(rec["סוגי רישיון"])),
        })
    df = pd.DataFrame(rows, columns=["תוצר", "יצרן", "מספר מופעים", "סוגי רישיון"])
    if not df.empty:
        df = df.sort_values("תוצר", kind="stable").reset_index(drop=True)
    return df


def main():
    if len(sys.argv) != 3:
        print(__doc__)
        sys.exit(1)
    src, out = Path(sys.argv[1]), Path(sys.argv[2])
    script_dir = Path(__file__).resolve().parent

    log_path = out.parent / "extraction.log"
    logging.basicConfig(filename=log_path, filemode="w", level=logging.INFO,
                        format="%(asctime)s %(levelname)s %(message)s", encoding="utf-8")

    try:
        table = load_translation(script_dir)
    except FileNotFoundError:
        print(f"חסר classifications.csv ליד הסקריפט ({script_dir})")
        sys.exit(1)

    files = [f for f in sorted(src.glob("*.docx")) if not f.name.startswith("~$")]
    if not files:
        print(f"לא נמצאו קבצי docx בתיקייה {src}")
        sys.exit(1)

    unk_path = out.parent / "unknown_classifications.csv"
    unk_path.unlink(missing_ok=True)  # מניעת קובץ ישן ומטעה מריצה קודמת

    buckets = {"יבואן ישיר": ([], []), "יבואן עקיף": ([], []), "יבואן זעיר": ([], [])}
    failed = 0
    for f in files:
        try:
            ltype, company, prods = extract_file(f, table)
            if ltype not in buckets:
                log.warning("%s: סוג רישיון לא זוהה — שויך לישיר לבדיקה", f.name)
                ltype = "יבואן ישיר"
            buckets[ltype][0].append(company)
            buckets[ltype][1].extend(prods)
            log.info("%s: %s, %d תוצרים, סטטוס=%s", f.name, ltype, len(prods),
                     company["סטטוס OCR"])
        except Exception:
            failed += 1
            log.exception("%s: כשל בחילוץ", f.name)

    with pd.ExcelWriter(out, engine="openpyxl", datetime_format="DD/MM/YYYY") as w:
        pd.DataFrame(buckets["יבואן ישיר"][0], columns=DIRECT_COMP_COLS
                     ).to_excel(w, sheet_name="ישיר-חברות", index=False)
        pd.DataFrame(buckets["יבואן ישיר"][1], columns=DIRECT_PROD_COLS
                     ).to_excel(w, sheet_name="ישיר-תוצרים", index=False)
        pd.DataFrame(buckets["יבואן עקיף"][0], columns=INDIRECT_COMP_COLS
                     ).to_excel(w, sheet_name="עקיף-חברות", index=False)
        pd.DataFrame(buckets["יבואן עקיף"][1], columns=INDIRECT_PROD_COLS
                     ).to_excel(w, sheet_name="עקיף-תוצרים", index=False)
        pd.DataFrame(buckets["יבואן זעיר"][0], columns=ZAIR_COMP_COLS
                     ).to_excel(w, sheet_name="זעיר-חברות", index=False)
        distinct_df = build_distinct_products(buckets)
        distinct_df.to_excel(w, sheet_name="תוצרים-DISTINCT", index=False)

    if unknown_class:
        with open(unk_path, "w", encoding="utf-8-sig", newline="") as f:
            wcsv = csv.writer(f)
            wcsv.writerow(["ערך", "מופעים", "קובץ לדוגמה"])
            for val, cnt in unknown_class.most_common():
                wcsv.writerow([val, cnt, unknown_example[val]])
        print(f"סיווגים לא מזוהים: {len(unknown_class)} -> {unk_path}")

    counts = {k: (len(v[0]), len(v[1])) for k, v in buckets.items()}
    print(f"הסתיים. ישיר: {counts['יבואן ישיר'][0]} חברות/{counts['יבואן ישיר'][1]} תוצרים | "
          f"עקיף: {counts['יבואן עקיף'][0]}/{counts['יבואן עקיף'][1]} | "
          f"זעיר: {counts['יבואן זעיר'][0]} חברות | תוצרים ייחודיים: {len(distinct_df)} | "
          f"כשלים: {failed} | לוג: {log_path}")


if __name__ == "__main__":
    main()
